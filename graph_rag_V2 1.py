import os
import uuid
import json
import time
from pathlib import Path
from datetime import datetime
from sqlite_mods_v2 import save_chat_to_db

#--------------- imported Ragas to evaluate------------------------------------------------------------

from ragas import evaluate                                                             
from ragas.metrics import (
    faithfulness,
    answer_relevancy,
    context_precision,
    context_recall,
   
)
from datasets import Dataset
from ragas.llms import LangchainLLMWrapper
from ragas.embeddings import LangchainEmbeddingsWrapper
# from ragas import EvaluationDataset
# from ragas.dataset_schema import SingleTurnSample

#-------------------------------------------------------------------------------------------------------

import pandas as pd
import streamlit as st
import pdfplumber
import docx
from langchain_core.messages import HumanMessage, AIMessage

from langchain_core.documents import Document
from langchain_experimental.text_splitter import SemanticChunker
from langchain_core.prompts import PromptTemplate
from langchain_classic.chains.question_answering import load_qa_chain
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_openai import ChatOpenAI

from neo4j import GraphDatabase


# ================================
# CONFIG
# ================================

UPLOAD_FOLDER = "documents"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Neo4j local
NEO4J_URI = "neo4j+s://11fac605.databases.neo4j.io"
NEO4J_USER = "11fac605"
NEO4J_PASSWORD = "c2oPkS8pU4iohNfBsLP5YBWHwv1WStK7kcvwtIi4CdU"

VECTOR_INDEX_NAME = "chunk_vector_index"


# ================================
# NEO4J DRIVER
# ================================

driver = GraphDatabase.driver(
    NEO4J_URI,
    auth=(NEO4J_USER, NEO4J_PASSWORD),
    max_connection_lifetime=3600,
    max_connection_pool_size=50,
    connection_timeout=30,
    keep_alive=True
)

# ================================
# EMBEDDING MODEL
# ================================

def get_embeddings():
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"}
    )

embedding_model = get_embeddings()

ragas_embeddings = LangchainEmbeddingsWrapper(embedding_model)      # -------------------------- adding embedding for RAGAS


# ================================
# LLM
# ================================

llm = ChatOpenAI(
    model="gpt-oss:20b",
    base_url="http://192.168.16.212:11434/v1",
    api_key="na",
    temperature=0.5,
    max_completion_tokens=1000
)

ragas_llm = LangchainLLMWrapper(llm)                            # adding llm for RAGAS  (giving access to one which we are using else we would have to add api)


# Cached file listing
@st.cache_data
def cached_list_files():
    return list_files()

# Cached document loading
@st.cache_data
def cached_load_file(f):
    return load_file(f)

# @st.cache_data
# def extract_question_entities_cached(question):
#     return extract_question_entities(question)

# ================================
# FILE MANAGEMENT
# ================================

def list_files():
    return list(Path(UPLOAD_FOLDER).glob("*"))


def save_uploads(uploaded):
    for file in uploaded:
        path = Path(UPLOAD_FOLDER) / file.name
        with open(path, "wb") as f:
            f.write(file.getbuffer())

def delete_files(files):
    for f in files:
        os.remove(f)


# ================================
# DOCUMENT LOADING
# ================================

def load_pdf(path: Path):

    documents = []

    try:
        with pdfplumber.open(str(path)) as pdf:

            for page_num, page in enumerate(pdf.pages, start=1):

                # -------- TEXT --------
                try:
                    text = page.extract_text()
                except Exception as e:
                    print(f"[WARN] Text extraction failed {path.name} page {page_num}: {e}")
                    text = None

                if text and text.strip():

                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "filename": path.name,
                                "page": page_num,
                                "type": "text",
                                "source": "pdf"
                            }
                        )
                    )

                # -------- TABLES --------
                try:
                    tables = page.extract_tables()
                except:
                    tables = []

                for table in tables:

                    if table:
                        try:

                            table_df = (
                                pd.DataFrame(table[1:], columns=table[0])
                                if len(table) > 1
                                else pd.DataFrame(table)
                            )

                            table_text = table_df.to_string(index=False)

                            documents.append(
                                Document(
                                    page_content=table_text,
                                    metadata={
                                        "filename": path.name,
                                        "page": page_num,
                                        "type": "table",
                                        "source": "pdf"
                                    }
                                )
                            )

                        except Exception as e:
                            print(f"[WARN] Table parse failed {path.name} page {page_num}: {e}")

    except Exception as e:
        print(f"[WARN] Skipping unreadable PDF {path.name}: {e}")
        return []

    return documents


def load_txt(path: Path):

    with open(path, "r", errors="ignore") as f:
        text = f.read()

    return [
        Document(
            page_content=text,
            metadata={
                "filename": path.name,
                "page": 1,
                "type": "text",
                "source": "txt"
            }
        )
    ]


def load_docx(path: Path):

    doc = docx.Document(path)

    text = "\n".join(
        [p.text for p in doc.paragraphs if p.text.strip()]
    )

    return [
        Document(
            page_content=text,
            metadata={
                "filename": path.name,
                "page": 1,
                "type": "text",
                "source": "docx"
            }
        )
    ]

def load_excel(path: Path):
    try:
        suffix = path.suffix.lower()

        if suffix == ".csv":
            df = pd.read_csv(path)
        elif suffix in [".xlsx", ".xls"]:
            df = pd.read_excel(path, engine="openpyxl")
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        text = df.to_string(index=False)

    except Exception as e:
        print(f"[WARN] Skipping file {path.name}: {e}")
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "filename": path.name,
                "page": 1,
                "type": "table",
                "source": "excel"
            }
        )
    ]

def load_file(path):

    # ext = path.suffix.lower()
    ext = Path(path.name).suffix.lower() if hasattr(path, "name") else Path(path).suffix.lower()

    if ext == ".pdf":
        return load_pdf(path)

    if ext == ".txt":
        return load_txt(path)

    if ext == ".docx":
        return load_docx(path)

    if ext in [".xlsx", ".xls", ".csv"]:
        return load_excel(path)

    return []


# ================================
# SEMANTIC CHUNKING
# ================================

def split_documents(documents):

    splitter = SemanticChunker(
        embedding_model,
        breakpoint_threshold_type="standard_deviation",
        breakpoint_threshold_amount=3.0,
    )

    return splitter.split_documents(documents)


def normalize_chunks(split_docs):

    chunks = []

    for doc in split_docs:

        topic = extract_main_topic(doc.page_content)

        chunks.append({
            "chunk_id": f"chunk_{uuid.uuid4().hex[:8]}",
            "chunk": doc.page_content,
            "summary": doc.page_content[:300],
            "main_topic": topic,
            "metadata": doc.metadata
        })

    return chunks


# ================================
# RELATION EXTRACTION
# ================================

SOURCE_CHUNK_PROMPT = """
You are a Graph RAG relationship extractor.

Allowed relation types:
BELONGS_TO_CATEGORY
RELATED_TO
APPLIES_TO
HAS_CONSEQUENCE
GOVERNS

Return JSON format:

{{
 "relations":[
  {{"type":"RELATION","target":"Entity"}}
 ]
}}
Main Topic:
{main_topic}

Chunk:
{chunk}
"""

def extract_relations(main_topic, chunk_text):

    prompt = SOURCE_CHUNK_PROMPT.format(
        main_topic=main_topic,
        chunk=chunk_text
    )

    response = llm.invoke(prompt)

    try:
        parsed = json.loads(response.content)
    except:
        return []

    # Handle both JSON formats
    if isinstance(parsed, dict):
        return parsed.get("relations", [])
    elif isinstance(parsed, list):
        return parsed
    else:
        return []


ENTITY_EXTRACTION_PROMPT = """
Extract important entities from the question.

Return JSON:
{{
  "entities": ["entity1", "entity2"]
}}

Question:
{question}
"""

def extract_question_entities(question):
    response = llm.invoke(
        ENTITY_EXTRACTION_PROMPT.format(question=question)
    )

    try:
        parsed = json.loads(response.content)
        return [e.lower().strip() for e in parsed.get("entities", [])]
    except:
        return []
    

MAIN_TOPIC_PROMPT = """
Extract the main topic in 1-3 words.

Chunk:
{chunk}
"""

def extract_main_topic(chunk):
    try:
        response = llm.invoke(
            MAIN_TOPIC_PROMPT.format(chunk=chunk)
        )
        return response.content.strip()
    except:
        return "general"
    

def graph_retrieval(question, selected_files=None, limit=50):

    # entities = extract_question_entities_cached(question)
    entities = extract_question_entities(question)

    # Debug (remove later if needed)
    print("Entities:", entities)

    if not entities:
        return []

    with driver.session() as session:

        # query = """                                              <------------------ Old query removed
        # MATCH (e:Entity)
        # WHERE any(ent IN $entities WHERE e.name CONTAINS ent)

        # MATCH (c:PolicyChunk)-[]->(e)

        # WHERE $files IS NULL OR c.filename IN $files

        # RETURN DISTINCT c.text AS text
        # LIMIT $limit
        # """                                                      <--------------------- removed

        # -----------------------------------------------removed version 2 --------------------------------------------------------------------------------
        # query = """                                                
        # MATCH (e:Entity)
        # WHERE any(ent IN $entities WHERE toLower(e.name) CONTAINS toLower(ent))

        # MATCH (c:PolicyChunk)-[]->(e)

        # WHERE $files IS NULL OR c.filename IN $files

        # WITH e.name AS entity, c.filename AS file, collect(DISTINCT c.text)[0..3] AS texts

        # WITH entity, collect({file: file, texts: texts}) AS file_data

        # WHERE size(file_data) > 1   

        # RETURN entity, file_data
        # LIMIT $limit  
        # """

        query = """
        MATCH (e:Entity)
        WHERE any(ent IN $entities WHERE toLower(e.name) = toLower(ent))   

        MATCH (c:PolicyChunk)-[]->(e)
        WHERE $files IS NULL OR c.filename IN $files

        WITH e.name AS entity, c.filename AS file, 
        collect(DISTINCT c.text)[0..50] AS texts

        WITH entity, collect({file: file, texts: texts}) AS file_data,
        count(DISTINCT file) AS file_count


        WITH entity, file_data, file_count
        ORDER BY file_count DESC, entity ASC   
        RETURN entity, file_data
        LIMIT $limit"""

        results = session.run(
            query,
            entities=entities,
            files=selected_files,
            limit=limit
        )

        return [r["text"] for r in results]
  

def vector_fallback(question, selected_files=None, k=5):

    if not selected_files:                                                                                                          # added to see if it makes the func work
        selected_files = None

    query_vector = list(map(float, embedding_model.embed_query(question)))
    
    #debug
    print("query vector length:  ",len(query_vector))
    # print("query vector ",query_vector[:10])

    with driver.session() as session:                                                                                        

        results = session.run(f"""
        CALL db.index.vector.queryNodes(
            '{VECTOR_INDEX_NAME}',
            $k,
            $vector
        )
        YIELD node, score
        WHERE $files IS NULL OR node.filename IN $files
        RETURN node.text AS text, node.filename AS file_name, score as score
        ORDER BY score DESC
        """,
        k=k,
        vector=query_vector,
        files=selected_files
        )
        
        # return [r["text"] for r in results]     


        return [
        {
            "text": r["text"],
            "file_name": r["file_name"],
            "score": r["score"]
        }
        for r in results
        ]                                                                          

    # with driver.session() as session:                                                                                                             REMOVED 

    #     results = session.run(f"""
    #     CALL db.index.vector.queryNodes(
    #     '{VECTOR_INDEX_NAME}',
    #     $k,
    #     $vector
    #     )
    #     YIELD node, score
    #     WHERE $files IS NULL OR size($files) = 0 OR node.filename IN $files
    #     RETURN 
    #     node.text AS text, 
    #     node.filename AS file_name,
    #     score AS score
    #     ORDER BY score DESC
    #     """,
    #     k=k,
    #     vector=query_vector,
    #     files=selected_files
    #     )

    #     results = [
    #     {
    #         "text": r["text"],
    #         "file_name": r["file_name"],
    #         "score": r["score"]
    #     }
    #     for r in results
    # ]

    # return results                                                                                                                                 REMOVED 
    
    # with driver.session() as session:
    #     results = session.run("""
    # CALL db.index.vector.queryNodes(
    #     'chunk_vector_index',
    #     $k,
    #     $vector
    # )
    # YIELD node, score
    # RETURN node.text AS text, score
    # ORDER BY score DESC
    # """,
    # k=5,
    # vector=query_vector
    # )

    #     rows = list(results)     # debug
    #     print("ROWS:", rows)                                                                                                                       REMOVED
       


# ================================
# GRAPH BUILDER
# ================================
def push_chunk(document, node, relations):

    ALLOWED_RELATIONS = {
        "BELONGS_TO_CATEGORY",
        "RELATED_TO",
        "APPLIES_TO",
        "HAS_CONSEQUENCE",
        "GOVERNS"
    }

    with driver.session() as session:

        # -------------------------
        # DOCUMENT NODE
        # -------------------------
        session.run("""
        MERGE (d:Document {id:$doc})
        SET d.filename=$file
        """,
        doc=document["id"],
        file=document["filename"]
        )

        # -------------------------
        # CHUNK NODE
        # -------------------------
        session.run("""
        MERGE (c:PolicyChunk {id:$id})
        SET c.text=$text,
            c.summary=$summary,
            c.main_topic=$topic,
            c.filename=$filename,
            c.page=$page,
            c.type=$type,
            c.source=$source
        """,
        id=node["id"],
        text=node["text"],
        summary=node["summary"],
        topic=node["main_topic"],
        filename=node["metadata"]["filename"],
        page=node["metadata"].get("page", 1),
        type=node["metadata"].get("type", "text"),
        source=node["metadata"].get("source", "unknown")
        )

        # -------------------------
        # DOCUMENT → CHUNK RELATION
        # -------------------------
        session.run("""
        MATCH (d:Document {id:$doc})
        MATCH (c:PolicyChunk {id:$chunk})
        MERGE (d)-[:HAS_CHUNK]->(c)
        """,
        doc=document["id"],
        chunk=node["id"]
        )

        # -------------------------
        # ENTITY + RELATIONS
        # -------------------------
        seen_entities = set()

        for rel in relations:

            # ---- VALIDATE RELATION ----
            rel_type = rel.get("type", "").strip()

            if rel_type not in ALLOWED_RELATIONS:
                continue  # skip invalid relation

            # ---- CLEAN ENTITY ----
            entity = rel.get("target", "").strip().lower()

            # remove junk words
            if entity in ["it", "this", "that", "they", "he", "she"]:
                continue

            if not entity or len(entity) < 3:
                continue  # skip junk like "a", "it"

            if entity in seen_entities:
                continue  # avoid duplicates

            seen_entities.add(entity)

            # ---- STORE ENTITY (NO EMBEDDING NEEDED) ----
            session.run(f"""
            MERGE (e:Entity {{name:$name}})
            WITH e
            MATCH (c:PolicyChunk {{id:$chunk}})
            MERGE (c)-[:{rel_type}]->(e)
            """,
            name=entity,
            chunk=node["id"]
            )

# ================================
# VECTOR STORAGE
# ================================

def store_embedding(chunk_id, text):

    # vector = embedding_model.embed_query(text)
    vector = list(map(float, embedding_model.embed_query(text)))

    with driver.session() as session:

        session.run("""
        MATCH (c:PolicyChunk {id:$id})
        SET c.embedding=$vec
        """,
        id=chunk_id,
        vec=vector
        )


# ================================
# GRAPH BUILD PIPELINE
# ================================

def build_graph_from_documents(documents, filename):

    split_docs = split_documents(documents)

    chunks = normalize_chunks(split_docs)

    for chunk in chunks:

        node = {
            "id": chunk["chunk_id"],
            "text": chunk["chunk"],
            "summary": chunk["summary"],
            "main_topic": chunk["main_topic"],
            "metadata": chunk["metadata"]
        }       

        relations = extract_relations(
            chunk["main_topic"],
            chunk["chunk"]
        )

        document = {
            "id": filename,
            "filename": filename
        }

        push_chunk(document, node, relations)

        store_embedding(node["id"], node["text"])


# ================================
# GRAPH RAG QUERY
# ================================
def graph_rag_query(question, top_k=5):   

    timings = {}
    context = []

    # -------------------------
    # STEP 1 — ENTITY EXTRACTION
    # -------------------------
    # t0 = time.time()

    # entities = extract_question_entities_cached(question)                                 ALREADY removed

    # entities = extract_question_entities(question)                                        Removed for RAGAS  -  DSB
    # timings["entity_extraction"] = round(time.time() - t0, 2)

    # Debug (remove later if needed)
    # print("Entities:", entities)                                                           Removed for RAGAS  -  DSB

    # -------------------------
    # STEP 2 — GRAPH RETRIEVAL
    # -------------------------
    # t1 = time.time()                                      ---------------- REMOVED FOR TESTING RAGAS ON VECTOR FALLBACK -------------------------

    # if entities:
    #     with driver.session() as session:

    #         results = session.run("""
    #         MATCH (e:Entity)
    #         WHERE any(ent IN $entities WHERE e.name CONTAINS ent)

    #         MATCH (c:PolicyChunk)-[r]->(e)

    #         OPTIONAL MATCH (c)-[:RELATED_TO|APPLIES_TO|HAS_CONSEQUENCE]->(e2)<-[:RELATED_TO|APPLIES_TO|HAS_CONSEQUENCE]-(other:PolicyChunk)

    #         RETURN DISTINCT c.text AS text
    #         LIMIT $limit
    #         """,
    #         entities=entities,
    #         limit=top_k
    #         )

    #         # for r in results:
    #         #     # context.append(r["text"])
    #         #     context = list(set(context))
    #         for r in results:
    #             context.append(r["text"])     
    #             # context.append({
    #             #     "text": r["text"],
    #             #     "file_name": r["file_name"] if r["file_name"] else "Unknown Source" 
    #             # })

    #         # deduplicate AFTER loop
    #         context = list(set(context))   
            
    # timings["graph_retrieval"] = round(time.time() - t1, 2)                    <-----------------------------REMOVED------------------------------------





    # -------------------------
    # STEP 3 — FALLBACK (VECTOR)
    # -------------------------
    t1 = time.time()

    # if not context:
    context = vector_fallback(question, k=top_k)  

    # print(" context after vector_fallback  ", context)                                                                     DEBUG                 

    # context = vector_fallback(question, k=top_k)
    # if context is None or len(context) == 0:
    #     return "No relevant information found in knowledge base.", timings, None

        

    timings["fallback"] = round(time.time() - t1, 2)

    # -------------------------
    # STEP 4 — LLM
    # -------------------------
    t2 = time.time()

    
    # context = context[:5]                                                        <-----was already removed

    # context_org = context[:5] if len(context) > 5 else context                          removed dsb                       #<----- original way
    # context_org = context[:5]                                                             removed dsb 

    # context = sorted(context, key=lambda x: x["score"], reverse=True)   < removed part 2
    # context_org = [c for c in context if c["score"] > 0.75][:5]

    # if len(context_org) == 0:
    #     context_org = context[:5]




    # --------------- removing for debugging----------------------------------------

    # if not context:
    #     context_org = []
    # else:
    #     context = sorted(context, key=lambda x: x["score"], reverse=True)

    #     context_org = [
    #         c for c in context
    #         if c["score"] > 0.75
    #     ]

    #     context_org = context_org[:5]

    # # fallback if filtering removes everything
    #     if len(context_org) == 0:
    #         context_org = context[:5]

    #--------------------------------------------------------------------------------------



    # v REMOVED BECAUSE IT WAS REDUNDANT ------------------------------------------------------------------------------Already sorted in cypher query (ORDER BY sort DESC) 

    # context = sorted(context, key=lambda x: x["score"], reverse=True)
    
    # context_org = [
    #     c for c in context
    #     # if c["score"] < 0.75]
    # ---------------------------------------------------------------------------------------------------------------

    context_org = context[:5]

    print("context: ",len(context))
    print("context_org: ",len(context_org))

   

#------------------------------------------------------RAGAS IMLPLEMENTATION---------------------------------------------------------------------------------

    
# FINAL FIX: ensure list of list
    # llm_context = context_org                                      # removed 
    llm_context = "\n\n".join([c["text"] for c in context_org])      # added for converting context into clean text format
    
    # ragas_context = [[c["text"] for c in context_org]]
    if len(context_org) == 0:
        ragas_context = [["No relevant context found"]]
    else:
        ragas_context = [[c["text"] for c in context_org]]


    print("llm context: ", llm_context)
    print("ragas context: ", ragas_context)


    # formatted_context = ""
    # for item in context:
    #     formatted_context += f"\n[Source File: {item['file_name']}]\nContent: {item['text']}\n"  

#------------------------------------------------------------------------------------------------------------------------------------------------

    prompt = f"""
        You are a GraphRAG assistant.
           
        Context:
        {llm_context}

        Question:
        {question}

        Answer clearly.
    """


    response = llm.invoke(prompt)

    timings["llm"] = round(time.time() - t2, 2)


    # --------------------------------------------------------------------------------------------------------------------------------------------------------------------
    # ---------------------------------------------ADDING RAGAS EVALUATION METRICS------------------------------------------------------------------------------------------

    ragas_scores = None

    try:                       

        eval_data = pd.DataFrame([{
            "question": str(question),
            "answer": str(response.content),
            "contexts": ragas_context[0] if isinstance(ragas_context[0], list) else [ragas_context[0]],
            "reference": (
                ragas_context[0][0] if isinstance(ragas_context[0], list)
                else ragas_context[0]
                )
        }])

        eval_dataset = Dataset.from_pandas(eval_data)
               

        ragas_result = evaluate(
            eval_dataset,
            metrics=[
                faithfulness,
                answer_relevancy,
                context_precision,
                context_recall,
               
            ],
            llm=ragas_llm,
            embeddings=ragas_embeddings


        )

        ragas_scores = ragas_result.to_pandas().to_dict(orient="records")[0]

    except Exception as e:
        ragas_scores = {"error": str(e)}

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#-------------------------------------------------------------------------------------------------------------------------------------------------------------------------






# -------------------------
#           TOTAL
# -------------------------
    timings["total"] = round(sum(timings.values()), 2)

    return response.content, timings, ragas_scores                                   #<--------------------- Returning RAGAS_scores

    

# ================================
#           STREAMLIT UI
# ================================
def FAQ():
    try:
        st.header("Graph RAG Assistant  (Practice) 💬💁")

        with st.form("input_form"):
            user_question = st.text_input("Ask question") 
            submitted = st.form_submit_button("Submit")

            if submitted and user_question.strip():

                try:
                    with st.spinner("Searching CRM knowledge base... 🔍"):

                        # ⏱ Start timer
                        start_time = time.time()

                        answer, timings, ragas_scores = graph_rag_query(user_question)

                        # ⏱ Total UI latency
                        latency = time.time() - start_time

                    if not answer:
                        st.warning("I couldn’t generate a response from the knowledge base.")
                        return

                    st.session_state.form_response = answer

                    # Streaming (no artificial delay)
                    def stream_data():
                        for word in answer.split(" "):
                            yield word + " "

                    if st.session_state.form_response:
                        st.markdown("**Reply:**")
                        st.write_stream(stream_data)

                        # ✅ Total time
                        st.caption(f"⏱ Total Response Time: {latency:.2f} sec")

                        # ✅ Detailed breakdown
                        with st.expander("⚙️ Performance Breakdown"):
                            st.json(timings)

                        with st.expander("📊 RAGAS Evaluation"):                       #---------------------- added expander for ragas score
                            st.json(ragas_scores)
                            
                    else:
                        st.markdown("Unable to get a valid response!")

                    # Save chat
                    if "user_id" in st.session_state:
                        save_chat_to_db(
                            st.session_state["user_id"],
                            st.session_state["username"],
                            [
                                HumanMessage(content=user_question),
                                AIMessage(content=answer)
                            ]
                        )

                except Exception as e:
                    st.error(f"Error in FAQ processing: {e}")
                    st.warning("Fallback response used.")

    except Exception as e:
        st.error(f"Error in FAQ: {e}")



# -----------------------------------------------------------COMPARE DOCUMENTS--------------------------------------------------------------------


def FAQ_compare_documents():
    st.header("Document QA / Comparison 📄🤝")

    files = cached_list_files()
    if not files:
        st.info("No documents available.")
        return

    # File selection
    with st.expander("Select Documents ▼", expanded=False):
        file_names = [f.name for f in files]
        selected_file_names = st.multiselect("Choose files", file_names, key="multi1")    #                   <----- key added for difference

    if not selected_file_names:
        st.warning("Please select at least one file.")
        return

    selected_files = [f for f in files if f.name in selected_file_names]

   # st.subheader("Ask your question:")                                         <----------- removed as we are hardcoding the question

    with st.form("qa_form"):

    #    user_question = st.text_input("Your question",key="question_compare")     <--------- ADDED KEY FOR differentitating between multiple questions    REMOVED

        user_question = "compare these documents and provide appropriate results"

        user_question2 = "provide summary of document in 1000 words"

        submitted = st.form_submit_button("Submit")

        if submitted and user_question.strip():

            with st.spinner("Searching... 🔍"):

                # CASE 1: SINGLE FILE
                
                if len(selected_files) == 1:

                    file_name = selected_files[0].name
                    context = []

                    # with driver.session() as session:
                    context = graph_retrieval(
                    user_question2,            
                    selected_files=[file_name],
                    limit=50 

                  )

                 # fallback if graph fails
                    if not context:
                        context = vector_fallback(
                            user_question2,              
                            selected_files=[file_name],
                            k = 50
                        )

                    if not context:
                        st.warning("No relevant answer found in this document.")
                        return

                    context_text = ""

                    context_text = "\n\n".join(context)

                    prompt = f"""
                    You are a strict document summarization assistant.

Task:
- Summarize the given document in approximately 1000 words.
- Provide a separate list of 8–12 key points directly from the document.
- Use ONLY the information present in the document.
- Do NOT guess or add any information not in the document.

Document:
{context_text}

User Question:
{user_question2}

Instructions:
1. Summary:
- Write a detailed, cohesive summary covering all major sections and ideas.
- Approximate length: 1000 words.
- If the document is shorter than 1000 words, summarize all content fully.
- Do not omit any major points.

2. Key Points:
- Extract 8–12 concise key points.
- Each point should be 1–2 sentences, clearly representing the main ideas.
- Reference content only from the document.

Output Format:

Summary:
detailed summary here

Key Points:
1. Key point 1
2. Key point 2
...
n. Key point n

If the document has no information to summarize, respond:
"Not found in document."
                    """
    #                 prompt = f"""                                                           <-----------------REMOVED old prompt
    # You are a strict assistant.

    # Answer ONLY from the given document.
    # Do NOT guess.

    # Document:
    # {context_text}

    # Question:
    # {user_question2}

    # If answer not found, say "Not found in document".                                      <-----------------REMOVED 
    # """
    
                #  CASE 2: MULTIPLE FILES
                
                else:

                    file_context = {}

                    # with driver.session() as session:
                        

                    for f in selected_files:
                        chunks = graph_retrieval(
                            user_question,
                            selected_files=[f.name],
                            limit=50
                        )
                        
                            # fallback
                        if not chunks:
                            chunks = vector_fallback(
                                user_question,
                                selected_files=[f.name],
                                k = 50
                            )

                        file_context[f.name] = chunks

                    context_text = ""

                    for file, chunks in file_context.items():
                        context_text += f"\n\n### Document: {file}\n"
                        context_text += "\n".join(chunks)
                    prompt = f"""
    You are a document comparison assistant.

    Compare the documents and answer the question.

    {context_text}

    Question:
    {user_question}

    Instructions:
    - Answer separately for each document
   
    - Then give comparison

    Format:

    Document:
    Answer:

    Comparison:
    """
    

                try:
                    answer = llm.invoke(prompt).content

                    st.markdown("**Reply:**")
                    st.write(answer)

                except Exception as e:
                    st.error(f"Error: {e}")

#----------------------------------------- ADDING FUNCTION FOR DIFFERENCE FINDER-----------------------------------------------------------------------
    
def FAQ_DIFF_FINDER():
    st.header("Document Difference Finder")

    files = cached_list_files()
    if not files:
        st.info("No documents available.")
        return

    # File selection
    with st.expander("Select Documents ▼", expanded=False):
        file_names = [f.name for f in files]
        selected_file_names = st.multiselect("Choose files", file_names, key="multi2")

    if not selected_file_names:
        st.warning("Please Select Atleast Two files.")
        return

    selected_files = [f for f in files if f.name in selected_file_names]

    st.subheader("Ask your question:")
    with st.form("qa_form2"):
        user_question = st.text_input("Your question",key="question_diff")
        submitted = st.form_submit_button("Submit")

        if submitted and user_question.strip():

            with st.spinner("Searching... 🔍"):

                # adding llm for paraphrasing the user_question--------------------------------------------------------------------------------------

                user_query_paraphrased_prompt =f"""
                you are a praphraser assistant, your job is to paraphrase the user question
                at least 10 questions which might produce best output from llm.
                user question:
                {user_question}
                provide output in form of list of paraphrased questions  """
                answer = llm.invoke(user_query_paraphrased_prompt).content

                st.write(answer)   # for debugging


                # CASE 1: SINGLE FILE                        <--------------                Removed single file option    
                
    #             if len(selected_files) == 1:

    #                 file_name = selected_files[0].name
    #                 context = []

    #                 # with driver.session() as session:
    #                 context = graph_retrieval(
    #                 user_question,
    #                 selected_files=[file_name],
    #                 limit=20
    #               )

    #              # fallback if graph fails
    #                 if not context:
    #                     context = vector_fallback(
    #                         user_question,
    #                         selected_files=[file_name]
    #                     )

    #                 if not context:
    #                     st.warning("No relevant answer found in this document.")
    #                     return

    #                 context_text = ""

    #                 context_text = "\n\n".join(context)

    #                 prompt = f"""
    # You are a strict assistant.

    # Answer ONLY from the given document and use ONLY given documents as source.
    # Do NOT guess.

    # Document:
    # {context_text}

    # Question:
    # {user_question}

    # If answer not found, say "Not found in document".                                                          < -----------   removed
    # """

    # ----------------------------------- CASE : MULTIPLE FILES------------------------------------------------------------------------------------------

                if len(selected_files) >= 2:
                    file_context = {}

                    # with driver.session() as session:
                        

                    for f in selected_files:
                        chunks = graph_retrieval(
                            user_question,
                            selected_files=[f.name],
                            limit=50                            #               <------------   increasing the number of chunks, to have more context
                        )
                        
                            # Vector fallback
                        if not chunks:
                            chunks = vector_fallback(
                                user_question,
                                selected_files=[f.name],
                                k = 50                            #               <------------   increasing the number of chunks, to have more context
                            )

                        file_context[f.name] = chunks

                    context_text = ""

                    for file, chunks in file_context.items():
                        context_text += f"\n\n### Document: {file}\n"
                        context_text += "\n".join(chunks)

#-------------------------------------------------------Removed Prompt--------------------Hardcoded 2 documents--------------------------------------------------------------
#                     prompt = f"""           
# You are a semantic document difference finder assistant.

# Your task is to identify and present meaningful differences between two documents.

# <context>
# {context_text}
# </context>

# User Question:
# {user_question}

# Instructions:

# 1. Determine mode of operation:
#    - If the User Question specifies a clear topic → ONLY compare documents on that topic.
#    - If the User Question is generic or does not specify a topic → follow the full process below.

# ----------------------------------------
# FULL COMPARISON PROCESS (for generic queries)
# ----------------------------------------

# Step 1: Topic Extraction
# - Read Document 1 and identify key topics (e.g., deadlines, pricing, responsibilities, scope, terms, names, etc.).
# - Topics should represent meaningful pieces of information, not generic labels.

# Step 2: Semantic Matching
# - For each extracted topic from Document 1:
#   - Search for the same or related information in Document 2 using semantic understanding (not exact keyword match).

# Step 3: Difference Identification
# - For each topic:
#   - Extract concise values from both documents.
#   - Identify whether the information is:
#     - Modified
#     - Added
#     - Removed
#     - Same (ignore if same)

# Step 4: Coverage Expansion
# - Repeat the same process by extracting additional topics to ensure all major aspects of the documents are covered.
# - Ensure no important section is skipped.

# Step 5: Internal Re-evaluation (CRITICAL)
# - Re-check the documents by reinterpreting the User Question in at least 2–3 different ways (e.g., "key changes", "major updates", "what differs").
# - Use this to catch any missed differences.
# - Do NOT show this step in output.

# ----------------------------------------

# Output format:

# | Topic | Document 1 | Document 2 | Difference |

# Rules:
# - "Topic" = clearly defined subject of difference
# - "Document 1" and "Document 2" = concise extracted values
# - "Difference":
#     - If both exist and differ: "Changed from [Doc1] to [Doc2]"
#     - If one is missing: "Present in Document X but not in Document Y"

# - DO NOT include rows where both documents are the same
# - Show multiple rows covering all major differences
# - Prioritize the most important differences first

# - If no meaningful differences are found, return:
# "No meaningful differences found."

# Do NOT include anything outside the table.

 
#     """-------------------------------------------------------------------removed------------------------------------------------------------------------------------------

                    prompt = f""" 
                      You are a semantic document difference finder assistant.

Your task is to identify and present meaningful differences across multiple documents.

<context>
{context_text}
</context>

User Question:
{answer}

Instructions:

1. Determine mode of operation:
   - If the User Question specifies a clear topic → ONLY compare documents on that topic.
   - If the User Question is generic or unclear → follow the full process below.

----------------------------------------
FULL COMPARISON PROCESS (for generic queries)
----------------------------------------

Step 1: Topic Extraction
- Read all documents and identify key topics.
- Topics should represent meaningful information (e.g., deadlines, pricing, responsibilities, scope, terms, names).

Step 2: Semantic Matching Across Documents
- For each topic:
  - Extract the corresponding value from EACH document using semantic understanding.
  - Do NOT rely on exact keyword matching.

Step 3: Difference Identification
- Compare values across all documents:
  - If ALL values are same → ignore
  - If values differ → include topic
  - If value exists in some documents but not others → mark as missing

Step 4: Coverage Expansion
- Ensure all major sections of documents are covered.
- Avoid redundant or trivial topics.

Step 5: Internal Re-evaluation (CRITICAL)
- Reinterpret the User Question in 2–3 different ways (e.g., "key differences", "major changes").
- Use this to catch missed differences.
- Do NOT show this step in output.

----------------------------------------

Output format:

| Topic | Doc 1 | Doc 2 | ... | Doc N | Difference |

Rules:
- Dynamically create one column per document (Doc 1, Doc 2, ..., Doc N), using each document’s actual name.
- "Topic" = clear subject of difference
- Each document cell = ULTRA-CONCISE value (max 5–10 words, no full sentences)

- "Difference":

- Explain differences between documents in a clear, context-aware way.
- Describe not just what changed, but how it affects meaning, intent, or implications.
- If values differ:
  - "Changed from [Doc1] to [Doc2], indicating [nature and impact of change]"
- If missing:
  - "Present in Document X but missing in Document Y, suggesting [implication]"
- For multiple documents: summarize patterns (e.g., alignment, deviation, progression).
- Focus on meaningful implications (e.g., cost, timeline, obligations, scope).
- Ensure the explanation helps the user understand why the change may have occurred.

- DO NOT include rows where all documents are the same
- Prioritize most important differences first
- Limit to most meaningful differences (avoid noise)

- If no meaningful differences are found, return:
"No meaningful differences found."

Do NOT include anything outside the table.
"""            
#-------------------------------------------------------------------Removed prompt (not working)------------------------------------------------------------------------------------------
#                     prompt = """
#                       You are a semantic document difference finder assistant.

# Your task is to identify and present meaningful differences across multiple documents.

# <context>
# {context_text}
# </context>

# User Question:
# {user_question}

# Instructions:

# 1. Determine mode of operation:
#    - If the User Question specifies a clear topic → ONLY compare documents on that topic.
#    - If the User Question is generic or unclear → follow the full process below.

# ----------------------------------------
# FULL COMPARISON PROCESS (for generic queries)
# ----------------------------------------

# Step 1: Topic Extraction (Adaptive)
# - Identify key topics from each document.
# - Topics should represent meaningful information (e.g., deadlines, pricing, responsibilities, scope, terms, names).
# - First assess overlap:
#   - If documents share comparable subject matter → extract aligned topics.
#   - If documents partially overlap → extract both shared topics AND document-specific topics.
#   - If little or no overlap → extract high-level comparison dimensions instead:
#     - Purpose
#     - Target Audience
#     - Document Type
#     - Key Intent
#     - Domain / Subject Area
#     - Structure / Format
# - Do NOT force alignment if topics do not naturally match.

# Step 2: Semantic Matching Across Documents
# - For each topic:
#   - Extract the corresponding value from EACH document using semantic understanding.
#   - If a topic is missing in a document, mark it explicitly.
#   - Do NOT rely on exact keyword matching.

# Step 3: Difference Identification
# - Include a topic if:
#   - Values differ across documents, OR
#   - Topic exists in some documents but not others (missing is a difference)
# - Ignore only if ALL documents have semantically identical values.

# Step 4: Coverage Expansion
# - Ensure all major sections OR high-level dimensions are covered.
# - Include both shared and unique topics when relevant.
# - Avoid redundant or trivial topics.

# Step 5: Internal Re-evaluation (CRITICAL)
# - Reinterpret the User Question in 2–3 different ways (e.g., "key differences", "major changes").
# - Use this to catch missed differences.
# - Do NOT show this step in output.

# ----------------------------------------

# Output format:

# | Topic | {Document Name 1} | {Document Name 2} | ... | {Document Name N} | Difference |

# Rules:
# - Dynamically create one column per document, using each document’s actual name.
# - "Topic" = clear subject of difference
# - Each document cell = ULTRA-CONCISE value (max 5–10 words, no full sentences)

# - "Difference":
#   - Explain differences in a clear, context-aware way in detail.
#   - Describe what changed AND how it affects meaning, intent, or implications.
#   - If values differ:
#     - "Changed from [{Document Name 1}] to [{Document Name 2}], indicating [nature and impact]"
#   - If missing:
#     - "Present in [{Document Name X}] but missing in [{Document Name Y}], suggesting [implication]"
#   - For multiple documents: summarize patterns (alignment, deviation, progression).
#   - If no shared topics exist, explain fundamental differences using high-level dimensions.
#   - Focus on meaningful implications (e.g., cost, timeline, obligations, scope, intent).

# - DO NOT include rows where all documents are the same
# - Prioritize most important differences first
# - Limit to most meaningful differences (avoid noise)

# - If no meaningful differences are found, return:
# "No meaningful differences found."

# Do NOT include anything outside the table.

# """                   
#----------------------------------------------------------------removed---------------------------------------------------------------------------------------------------------                
                
                else:
                    st.write("ATLEAST 2 DOCUMENTS")
                    

                    
                # =========================
                # LLM CALL
                # =========================
                try:
                    answer = llm.invoke(prompt).content

                    st.markdown("**Reply:**")
                    st.write(answer)

                except Exception as e:
                    st.error(f"Error: {e}")



                
# ================================
# DOCUMENT INGESTION
# ================================

def process_all_files():

    for f in list_files():

        docs = load_file(f)

        build_graph_from_documents(docs, f.name)


# ================================
# MAIN APP
# ================================
def main():

    st.title("Graph RAG Knowledge Assistant")
#     # File uploader
    uploaded = st.file_uploader(
        "Upload documents",
        accept_multiple_files=True
    )

    if uploaded:
        with st.spinner("Uploading files and building knowledge graph... 🔄"):
            # Save uploaded files
            save_uploads(uploaded)

            # Process uploaded files immediately
            for f in uploaded:
                docs = load_file(f)
                build_graph_from_documents(docs, f.name)

        st.success("Files uploaded and knowledge graph updated successfully ✅")

    FAQ()
   
    st.divider()    # visual separation
    FAQ_compare_documents()
    st.divider()
    FAQ_DIFF_FINDER()

if __name__ == "__main__":
    main()